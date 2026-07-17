"""Document renderer — Markdown always, DOCX/PDF if toolchain available."""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from .schemas import RenderManifest


def _hash_file(path: Path) -> str:
    if not path.exists():
        return ""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _docx_available() -> tuple[bool, str]:
    try:
        import docx  # noqa: F401
        return True, ""
    except Exception:
        if shutil.which("pandoc"):
            return True, ""
        return False, "python-docx not installed and pandoc not on PATH"


def _pdf_available() -> tuple[bool, str]:
    if shutil.which("pandoc"):
        return True, ""
    if shutil.which("libreoffice") or shutil.which("soffice"):
        return True, ""
    try:
        import reportlab  # noqa: F401
        return True, ""
    except Exception:
        return False, "no pandoc/libreoffice/reportlab available"


def render_document(document_id: str, markdown_text: str, output_dir: str) -> RenderManifest:
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    md_path = out / f"{document_id}.md"
    md_path.write_text(markdown_text, encoding="utf-8")

    manifest = RenderManifest(
        document_id=document_id,
        source_markdown=str(md_path),
        md_path=str(md_path),
        render_attempted=True,
        md_rendered=True,
        source_hash=hashlib.sha256(markdown_text.encode("utf-8")).hexdigest(),
        md_hash=_hash_file(md_path),
    )

    docx_ok, docx_reason = _docx_available()
    manifest.toolchain_available["docx"] = docx_ok
    if docx_ok:
        try:
            docx_path = out / f"{document_id}.docx"
            _render_docx(markdown_text, docx_path)
            manifest.docx_path = str(docx_path)
            manifest.docx_rendered = docx_path.exists()
            manifest.docx_hash = _hash_file(docx_path)
            manifest.toolchain_used["docx"] = "python-docx"
        except Exception as e:
            manifest.toolchain_available["docx"] = False
            manifest.toolchain_missing_reasons["docx"] = f"render error: {e}"
    else:
        manifest.toolchain_missing_reasons["docx"] = docx_reason

    pdf_ok, pdf_reason = _pdf_available()
    manifest.toolchain_available["pdf"] = pdf_ok
    if pdf_ok:
        try:
            pdf_path = out / f"{document_id}.pdf"
            rendered = _render_pdf(markdown_text, pdf_path)
            manifest.pdf_path = str(pdf_path) if rendered else ""
            manifest.pdf_rendered = rendered and pdf_path.exists()
            manifest.pdf_hash = _hash_file(pdf_path) if rendered else ""
            if not rendered:
                manifest.toolchain_missing_reasons["pdf"] = "renderer present but produced no file"
        except Exception as e:
            manifest.toolchain_available["pdf"] = False
            manifest.toolchain_missing_reasons["pdf"] = f"render error: {e}"
    else:
        manifest.toolchain_missing_reasons["pdf"] = pdf_reason

    return manifest


def _render_docx(markdown_text: str, docx_path: Path) -> None:
    import docx
    doc = docx.Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(str(docx_path))


def _render_pdf(markdown_text: str, pdf_path: Path) -> bool:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        y = 750
        for line in markdown_text.splitlines():
            if y < 50:
                c.showPage()
                y = 750
            c.drawString(50, y, line[:100])
            y -= 14
        c.save()
        return pdf_path.exists()
    except Exception:
        return False
