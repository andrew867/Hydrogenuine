"""Pack 12: DOCX builder — create, add heading/paragraph/table, finalize to file. Tenant-scoped path."""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from hg_core.docs.paths import get_exports_root


_docx_buffers: Dict[str, Any] = {}


def docx_create(title: str, tenant_id: str) -> str:
    """Create a new DOCX buffer. Returns doc_id (internal id for add_* / finalize)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx required; pip install python-docx")
    doc_id = str(uuid.uuid4())
    doc = Document()
    doc.add_heading(title, level=0)
    _docx_buffers[doc_id] = {"doc": doc, "tenant_id": tenant_id}
    return doc_id


def docx_add_heading(doc_id: str, text: str, level: int = 1) -> None:
    doc_buf = _docx_buffers.get(doc_id)
    if not doc_buf:
        raise ValueError("doc_id not found")
    doc_buf["doc"].add_heading(text, level=level)


def docx_add_paragraph(doc_id: str, text: str, citations: Optional[List[Dict[str, Any]]] = None) -> None:
    doc_buf = _docx_buffers.get(doc_id)
    if not doc_buf:
        raise ValueError("doc_id not found")
    p = doc_buf["doc"].add_paragraph(text)
    if citations:
        for c in citations:
            ref = f"[{c.get('document_id', '')} p{c.get('page_start', '')}-{c.get('page_end', '')}]"
            doc_buf["doc"].add_paragraph(ref, style="Intense Quote")


def docx_add_table(doc_id: str, headers: List[str], rows: List[List[str]]) -> None:
    doc_buf = _docx_buffers.get(doc_id)
    if not doc_buf:
        raise ValueError("doc_id not found")
    table = doc_buf["doc"].add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            if c_idx < len(table.rows[r_idx + 1].cells):
                table.rows[r_idx + 1].cells[c_idx].text = str(cell)
    doc_buf["doc"].add_paragraph()


def docx_finalize(doc_id: str, filename: str) -> str:
    """Write DOCX to tenant exports dir. Returns file_id for GET /v1/files/{file_id}/download."""
    doc_buf = _docx_buffers.pop(doc_id, None)
    if not doc_buf:
        raise ValueError("doc_id not found")
    tenant_id = doc_buf["tenant_id"]
    export_root = get_exports_root(tenant_id)
    export_root.mkdir(parents=True, exist_ok=True)
    file_id = str(uuid.uuid4())
    safe_name = (filename or "export").strip().replace("..", "") or "export"
    if not safe_name.lower().endswith(".docx"):
        safe_name += ".docx"
    path = export_root / f"{file_id}.docx"
    doc_buf["doc"].save(str(path))
    return file_id
